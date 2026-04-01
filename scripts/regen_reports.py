#!/usr/bin/env python3
"""
regen_reports.py
================
기존 classified CSV + JSON에서 xlsx / docx / md 재생성
(LLM 재호출 없음 — 결정론적 재생성)

사용법:
    python scripts/regen_reports.py 2026-03-31
"""

import json, os, sys
import pandas as pd
from datetime import datetime
from dateutil.parser import parse as dateparse

# ── 날짜 설정 ──
if len(sys.argv) > 1:
    TARGET_DATE = dateparse(sys.argv[1]).date()
else:
    print("사용법: python scripts/regen_reports.py YYYY-MM-DD")
    sys.exit(1)

DATE_TAG  = TARGET_DATE.strftime('%Y%m%d')
MONITOR_DIR = 'monitoring'
DAILY_DIR   = os.path.join(MONITOR_DIR, DATE_TAG)
LLM_REPORT_MODEL = "claude-sonnet-4-6"

print(f"{'='*60}")
print(f"  KMI 리포트 재생성: {TARGET_DATE}")
print(f"{'='*60}\n")

# ── 카테고리 정의 ──
CAT_KR = {
    "1_Security":       "안보·군사",
    "2_Safety":         "해양안전",
    "3_Freight":        "운임·물류비",
    "4_PortCargo":      "항만·화물",
    "5_EconFinance":    "경제·금융",
    "6_Seafood":        "수산물",
    "7_Shipping":       "해운",
    "8_Logistics":      "물류",
    "9_PortCongestion": "항만혼잡",
    "10_OtherIndustry": "기타 산업",
}
CAT_ORDER = list(CAT_KR.keys())

TRACKED_KEYWORDS = [
    ('호르무즈|Hormuz',                '호르무즈'),
    ('이란|Iran|IRGC',                 '이란'),
    ('트럼프|Trump',                   '트럼프'),
    ('나프타|naphtha|납사',             '나프타'),
    ('LNG',                            'LNG'),
    ('원유|crude oil|petroleum',       '원유'),
    ('유가|oil price',                 '유가'),
    ('코스피|KOSPI',                   '코스피'),
    ('환율|exchange rate|won-dollar',  '환율'),
    ('수에즈|Suez',                    '수에즈'),
    ('말라카|Malacca',                 '말라카'),
    ('운임|freight rate|tanker rate',  '운임'),
    ('셧다운|shutdown',                '셧다운'),
    ('비상경영|emergency',              '비상경영'),
    ('OECD',                           'OECD'),
    ('반도체|semiconductor',           '반도체'),
    ('석유화학|petrochemical',          '석유화학'),
    ('물가|inflation|CPI',             '물가'),
    ('홍해|Red Sea|Houthi',           '홍해'),
    ('해운|shipping|carrier',          '해운'),
]

# ══════════════════════════════════════════════════════════════
# 1. 파일 존재 확인
# ══════════════════════════════════════════════════════════════
gdelt_csv  = os.path.join(DAILY_DIR, f'gdelt_mon_classified_daily_{DATE_TAG}.csv')
naver_csv  = os.path.join(DAILY_DIR, f'naver_mon_classified_daily_{DATE_TAG}.csv')
json_path  = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.json')

for f in [gdelt_csv, naver_csv, json_path]:
    if not os.path.exists(f):
        print(f"❌ 파일 없음: {f}")
        sys.exit(1)
    print(f"✅ 확인: {f}")

# ── 로드 ──
gdelt_cls = pd.read_csv(gdelt_csv, encoding='utf-8-sig')
naver_cls = pd.read_csv(naver_csv, encoding='utf-8-sig')
with open(json_path, 'r', encoding='utf-8') as f:
    result_data = json.load(f)
llm = result_data.get('llm_result', {})

print(f"\nGDELT: {len(gdelt_cls)}건  네이버: {len(naver_cls)}건\n")

# ══════════════════════════════════════════════════════════════
# 2. Excel 재생성
# ══════════════════════════════════════════════════════════════
print("[1/3] Excel 생성...")
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    dfs = []
    for df, label in [(gdelt_cls, 'GDELT'), (naver_cls, '네이버')]:
        df = df.copy()
        df['source'] = label
        dfs.append(df)
    combined = pd.concat(dfs, ignore_index=True)

    xlsx_path = os.path.join(DAILY_DIR, f'daily_report_{DATE_TAG}.xlsx')
    wb = openpyxl.Workbook()

    header_fill = PatternFill('solid', fgColor='1F4E79')
    header_font = Font(bold=True, color='FFFFFF')
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── 시트1: 요약 ──
    ws_sum = wb.active
    ws_sum.title = '요약'
    ws_sum['A1'] = f'KMI 해상공급망 일일 모니터링 — {TARGET_DATE}'
    ws_sum['A1'].font = Font(bold=True, size=14)

    row = 3
    ws_sum.cell(row, 1, '구분').font = header_font
    ws_sum.cell(row, 1).fill = header_fill
    for col, lbl in enumerate(['HIGH', 'MEDIUM', 'LOW', 'NONE', '합계'], 2):
        ws_sum.cell(row, col, lbl).font = header_font
        ws_sum.cell(row, col).fill = header_fill

    row += 1
    for src in ['GDELT', '네이버', '합계']:
        sub = combined if src == '합계' else combined[combined['source'] == src]
        ws_sum.cell(row, 1, src)
        for col, rel in enumerate(['HIGH', 'MEDIUM', 'LOW', 'NONE'], 2):
            ws_sum.cell(row, col, len(sub[sub['relevance'] == rel]))
        ws_sum.cell(row, 6, len(sub))
        row += 1

    # ── 시트2: 카테고리별 ──
    ws_cat = wb.create_sheet('카테고리별')
    ws_cat['A1'] = '카테고리별 HIGH+MEDIUM 기사 수'
    ws_cat['A1'].font = Font(bold=True, size=12)
    hm = combined[combined['relevance'].isin(['HIGH', 'MEDIUM'])]
    row = 3
    for cat in CAT_ORDER:
        kr = CAT_KR.get(cat, cat)
        n = len(hm[hm['category'] == cat])
        ws_cat.cell(row, 1, kr)
        ws_cat.cell(row, 2, n)
        row += 1

    # ── 시트3: 추적 키워드 ──
    ws_kw = wb.create_sheet('추적키워드')
    ws_kw['A1'] = '주요 키워드 언급 현황'
    ws_kw['A1'].font = Font(bold=True, size=12)
    row = 3
    for pattern_str, label in TRACKED_KEYWORDS:
        mask = combined['title'].str.contains(pattern_str, case=False, na=False, regex=True)
        ws_kw.cell(row, 1, label)
        ws_kw.cell(row, 2, int(mask.sum()))
        row += 1

    # ── 시트4: HIGH 기사 목록 ──
    ws_high = wb.create_sheet('HIGH기사')
    high_df = combined[combined['relevance'] == 'HIGH'][
        ['source', 'title', 'category', 'query_keyword', 'url']
    ].head(100)
    for col_idx, col_name in enumerate(high_df.columns, 1):
        ws_high.cell(1, col_idx, col_name).font = header_font
        ws_high.cell(1, col_idx).fill = header_fill
    for r_idx, row_data in enumerate(high_df.itertuples(index=False), 2):
        for c_idx, val in enumerate(row_data, 1):
            ws_high.cell(r_idx, c_idx, str(val) if val else '')

    wb.save(xlsx_path)
    print(f"  ✅ Excel 저장: {xlsx_path}")
except Exception as e:
    print(f"  ❌ Excel 실패: {e}")

# ══════════════════════════════════════════════════════════════
# 3. Markdown 재생성 (v3 notebook 동일 구조)
# ══════════════════════════════════════════════════════════════
print("[2/3] Markdown 생성...")

# ── 통계 준비 ──
dfs_all = []
for df, label in [(gdelt_cls, 'GDELT'), (naver_cls, '네이버')]:
    _df = df.copy()
    _df['source'] = label
    dfs_all.append(_df)
combined_all = pd.concat(dfs_all, ignore_index=True) if dfs_all else pd.DataFrame()
combined_all['source_type'] = combined_all.get('sourcecountry', pd.Series(dtype=str)).apply(
    lambda x: 'domestic' if x == 'South Korea' else 'international'
)
total_all   = len(combined_all)
hm_all      = combined_all[combined_all['relevance'].isin(['HIGH','MEDIUM'])]
hm_intl_all = hm_all[hm_all['source_type'] == 'international']
hm_dom_all  = hm_all[hm_all['source_type'] == 'domestic']

cat_stats = {}
for cat in CAT_ORDER:
    cat_hm = hm_all[hm_all['category'] == cat]
    h = len(cat_hm[cat_hm['relevance'] == 'HIGH'])
    m = len(cat_hm[cat_hm['relevance'] == 'MEDIUM'])
    cat_stats[cat] = {'high': h, 'med': m, 'total': h + m}

try:
    md_path = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.md')
    lines = []
    L = lines.append
    L(f"# 해상 공급망 모니터링 일일 분석 — {TARGET_DATE}")
    L(f"")
    L(f"> 생성 모델: {LLM_REPORT_MODEL}  |  분석 기준: HIGH+MEDIUM {len(hm_all)}건 "
      f"(해외 {len(hm_intl_all)}건 / 국내 {len(hm_dom_all)}건)")
    L(f"")
    L("---")
    L("")
    L("## 📌 오늘의 핵심")
    L("")
    L(llm.get("executive_summary", ""))
    L("")
    L("---")
    L("")
    L("## 카테고리별 분석")
    L("")
    cats_out = llm.get("categories", {})
    for cat in CAT_ORDER:
        if cat not in cats_out: continue
        cat_data = cats_out[cat]
        cat_kr   = CAT_KR.get(cat, cat)
        s        = cat_stats.get(cat, {})
        h, m     = s.get('high', 0), s.get('med', 0)
        if h + m == 0: continue
        L(f"### {cat_kr}  _(HIGH {h} / MEDIUM {m})_")
        L("")
        overseas     = cat_data.get("overseas", "").strip() if isinstance(cat_data, dict) else ""
        korea_impact = cat_data.get("korea_impact", "").strip() if isinstance(cat_data, dict) else ""
        if overseas:
            L("**🌐 해외 상황**"); L(""); L(overseas); L("")
        if korea_impact:
            L("**🇰🇷 국내 영향**"); L(""); L(korea_impact); L("")
    L("---"); L("")
    L("## 통계 요약"); L("")
    L("| 구분 | 건수 | 비중 |"); L("|------|-----:|-----:|")
    n_high = int(result_data.get('n_high', 0))
    n_med  = int(result_data.get('n_med', 0))
    n_low  = int(result_data.get('n_low', 0))
    n_none = int(result_data.get('n_none', 0))
    n_total = int(result_data.get('n_total', total_all))
    for label, cnt in [('총 수집', n_total), ('HIGH', n_high),
                       ('MEDIUM', n_med), ('LOW', n_low), ('NONE', n_none)]:
        pct = cnt/n_total*100 if n_total else 0
        L(f"| {label} | {cnt:,} | {pct:.1f}% |")
    L("")
    L("| 카테고리 | HIGH | MEDIUM | 합계 |"); L("|----------|-----:|-------:|-----:|")
    for cat in CAT_ORDER:
        s = cat_stats.get(cat, {})
        if s.get('total', 0) > 0:
            L(f"| {CAT_KR[cat]} | {s['high']} | {s['med']} | {s['total']} |")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"  ✅ MD 저장: {md_path}")
except Exception as e:
    print(f"  ❌ MD 실패: {e}")

# ══════════════════════════════════════════════════════════════
# 4. DOCX 재생성 (v3 notebook 동일 구조)
# ══════════════════════════════════════════════════════════════
print("[3/3] DOCX 생성...")
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0); sec.right_margin  = Cm(3.0)
    p = doc.add_heading(f'해상 공급망 모니터링 일일 분석 — {TARGET_DATE}', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f'생성 모델: {LLM_REPORT_MODEL}  |  HIGH+MEDIUM {len(hm_all)}건 '
        f'(해외 {len(hm_intl_all)}건 / 국내 {len(hm_dom_all)}건)'
    )
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()
    doc.add_heading('오늘의 핵심', level=2)
    doc.add_paragraph(llm.get('executive_summary', ''))
    doc.add_paragraph()
    doc.add_heading('카테고리별 분석', level=2)
    cats_out = llm.get('categories', {})
    for cat in CAT_ORDER:
        if cat not in cats_out: continue
        cat_data = cats_out[cat]
        cat_kr   = CAT_KR.get(cat, cat)
        s        = cat_stats.get(cat, {})
        h, m_cnt = s.get('high', 0), s.get('med', 0)
        if h + m_cnt == 0: continue
        doc.add_heading(f'{cat_kr}  (HIGH {h} / MEDIUM {m_cnt})', level=3)
        overseas     = cat_data.get('overseas', '').strip() if isinstance(cat_data, dict) else ''
        korea_impact = cat_data.get('korea_impact', '').strip() if isinstance(cat_data, dict) else ''
        if overseas:
            p = doc.add_paragraph(); p.add_run('🌐 해외 상황').bold = True
            doc.add_paragraph(overseas)
        if korea_impact:
            p = doc.add_paragraph(); p.add_run('🇰🇷 국내 영향').bold = True
            doc.add_paragraph(korea_impact)
    doc.add_page_break()
    doc.add_heading('통계 요약', level=2)
    tbl = doc.add_table(rows=6, cols=3); tbl.style = 'Table Grid'
    for j, h_txt in enumerate(['구분', '건수', '비중']):
        tbl.rows[0].cells[j].text = h_txt
        tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i_r, (label, cnt) in enumerate([
        ('총 수집', n_total), ('HIGH', n_high),
        ('MEDIUM', n_med), ('LOW', n_low), ('NONE', n_none)
    ]):
        tbl.rows[i_r+1].cells[0].text = label
        tbl.rows[i_r+1].cells[1].text = str(cnt)
        pct = cnt/n_total*100 if n_total else 0
        tbl.rows[i_r+1].cells[2].text = f"{pct:.1f}%"
    docx_path = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.docx')
    doc.save(docx_path)
    print(f"  ✅ DOCX 저장: {docx_path}")
except Exception as e:
    print(f"  ❌ DOCX 실패: {e}")

print(f"\n{'='*60}")
print(f"  재생성 완료: monitoring/{DATE_TAG}/")
print(f"{'='*60}")
