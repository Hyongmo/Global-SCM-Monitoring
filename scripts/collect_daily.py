#!/usr/bin/env python3
"""
collect_daily.py
================
해상 공급망 위기 일일 모니터링 자동화 스크립트
(gdelt_news_monitoring_v3.ipynb 를 GitHub Actions 용 스크립트로 변환)

실행 방식:
    python scripts/collect_daily.py [YYYY-MM-DD]
    날짜 미지정 시 어제(yesterday) 자동 사용

환경변수:
    ANTHROPIC_API_KEY    Claude API 키
    NAVER_CLIENT_ID      네이버 검색 API ID
    NAVER_CLIENT_SECRET  네이버 검색 API 시크릿
"""

import json, os, sys, hashlib, time, re, html, glob, requests
import pandas as pd
import networkx as nx
from collections import Counter
from datetime import datetime, timedelta, date
from dateutil.parser import parse as dateparse
from gdeltdoc import GdeltDoc, Filters
import anthropic

# ══════════════════════════════════════════════════════════════
# 0. 기본 설정
# ══════════════════════════════════════════════════════════════

# ── 수집 날짜 결정 (CLI 인수 또는 어제) ──
if len(sys.argv) > 1:
    TARGET_DATE = dateparse(sys.argv[1]).date()
else:
    TARGET_DATE = (datetime.now() - timedelta(days=1)).date()

d_start = TARGET_DATE
d_end   = TARGET_DATE
collect_dates   = [TARGET_DATE]
dates_to_collect = [TARGET_DATE]

DATE_TAG = TARGET_DATE.strftime('%Y%m%d')

print(f"{'='*60}")
print(f"  KMI 해상 공급망 일일 모니터링")
print(f"  수집 날짜: {TARGET_DATE}")
print(f"{'='*60}\n")

# ── 경로 설정 ──
MONITOR_DIR = 'monitoring'
CUMUL_DIR   = os.path.join(MONITOR_DIR, 'cumulative')
WEEKLY_DIR  = os.path.join(MONITOR_DIR, 'weekly')
DAILY_DIR   = os.path.join(MONITOR_DIR, DATE_TAG)

for d in [MONITOR_DIR, CUMUL_DIR, WEEKLY_DIR, DAILY_DIR]:
    os.makedirs(d, exist_ok=True)

# ── 쿼리 파일 ──
MON_QUERY_FILE = 'news_queries_monitoring_v2.json'

# ── API 클라이언트 ──
gd = GdeltDoc()
client = anthropic.Anthropic()  # ANTHROPIC_API_KEY 자동 사용

# ── 테스트 모드 (환경변수 TEST_SCALE 로 스케일 축소 가능) ──
_test_scale = int(os.environ.get('TEST_SCALE', 0))  # 0=정상, 양수=샘플 수 제한

# ── GDELT 파라미터 ──
MAX_RECORDS    = 10 if _test_scale else 250
SLEEP_SEC      = 0.5
LANGUAGES      = 'English'
MIN_KEYWORD_LEN = 5
MAX_RETRIES    = 2

# ── 네이버 파라미터 ──
NAVER_CLIENT_ID     = os.environ.get('NAVER_CLIENT_ID', '')
NAVER_CLIENT_SECRET = os.environ.get('NAVER_CLIENT_SECRET', '')
NAVER_API_URL = 'https://openapi.naver.com/v1/search/news.json'
NAVER_DISPLAY = 30 if _test_scale else 100
NAVER_SLEEP   = 0.3

# ── LLM 파라미터 ──
MODEL      = "claude-haiku-4-5-20251001"
BATCH_SIZE = 20
MAX_LLM_SAMPLE_GDELT = _test_scale if _test_scale else 1500
MAX_LLM_SAMPLE_NAVER = _test_scale if _test_scale else 1500
MIN_PER_KEYWORD      = 1 if _test_scale else 3

if _test_scale:
    print(f"⚠ TEST MODE: scale={_test_scale} (rec/kw={MAX_RECORDS}, "
          f"kw=도메인별2개, sample={MAX_LLM_SAMPLE_GDELT}/{MAX_LLM_SAMPLE_NAVER})")

# ── 카테고리 정의 ──
CATEGORIES = {
    "1_Security":       "Security — military threats, geopolitical tensions, chokepoint control, naval operations",
    "2_Safety":         "Safety — vessel accidents, maritime safety incidents, IMO regulations",
    "3_Freight":        "Freight Rates / Transit Volume — freight index changes (BDI, SCFI, etc.), traffic volume shifts",
    "4_PortCargo":      "Port Cargo Volume — import cargo volume, port throughput, port conditions",
    "5_EconFinance":    "Economy / Finance — stock prices, exchange rates, financial index movements, economic indicators",
    "6_Seafood":        "Seafood Trade — export/import value and unit prices of seafood, fishing industry impacts",
    "7_Shipping":       "Shipping Industry — domestic/overseas shipping companies, carrier trends, alternative routes, fleet changes",
    "8_Logistics":      "Logistics Companies — impacts on logistics/forwarding companies, response measures",
    "9_PortCongestion": "Port Congestion — port delays, berth waiting, container dwell time",
    "10_OtherIndustry": "Other Industries — petroleum/naphtha, LNG, fertilizers, bunkering price changes, petrochemical impacts",
}
CATEGORY_KEYS = list(CATEGORIES.keys())

CAT_KR = {
    "1_Security":       "안보·군사",
    "2_Safety":         "해양안전",
    "3_Freight":        "운임·물류비",
    "4_PortCargo":      "항만·화물",
    "5_EconFinance":    "경제·금융",
    "6_Seafood":        "수산물",
    "7_Shipping":       "해운",
    "8_Logistics":      "물류",
    "9_PortCongestion": "항만혼잡",
    "10_OtherIndustry": "기타 산업",
}
CAT_ORDER = list(CAT_KR.keys())

CAT_BLOCK_MON = """  - 1_Security: International relations changes — transit fee disputes, military deployments, UN resolutions,
    joint condemnation statements, naval operations, chokepoint control, Iran tensions, Houthi attacks
  - 2_Safety: Vessel accidents, casualties, IMO emergency meetings, maritime safety regulations,
    Iran communications to IMO, safety incidents at sea
  - 3_Freight: Freight index changes (BDI, SCFI, CCFI, WCI), traffic volume shifts through Suez/Hormuz/Cape/Aden,
    tanker rate spikes, container freight surcharges, shipping rate trends
  - 4_PortCargo: Import cargo volume changes (crude oil, LNG, petroleum products, refined products),
    port throughput data, import port status, berth utilization
  - 5_EconFinance: Stock prices (KOSPI, MSCI), VIX, exchange rates (KRW/USD), Brent/WTI oil price movements,
    financial market reactions in US/Iran/Korea, economic indicators
  - 6_Seafood: Seafood export/import values and unit prices, fishing industry difficulties from oil price rises,
    coastal/deep-sea fishing company impacts, seafood market inflation
  - 7_Shipping: Domestic shipping company (HMM, Pan Ocean) trends, top 10 global carrier updates (Maersk, MSC, CMA CGM),
    alternative route strategies, emergency bunker surcharges (EBS), fleet repositioning
  - 8_Logistics: Logistics company impacts (CJ Logistics, Hyundai Glovis), logistics cost changes,
    operational difficulties, shipper service disruptions, forwarding company responses
  - 9_PortCongestion: Port congestion around Hormuz/Gulf region, cargo throughput delays,
    berth waiting times, container dwell time increases, port capacity constraints
  - 10_OtherIndustry: Petroleum/naphtha price changes, LNG spot price spikes, fertilizer price impacts,
    bunkering price changes, petrochemical feedstock costs, refinery margin shifts"""

# ══════════════════════════════════════════════════════════════
# 0-b. KG 로드 + 엔티티 매칭 (seed_kg_v4.json)
# ══════════════════════════════════════════════════════════════

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT  = os.path.dirname(_SCRIPT_DIR)
KG_PATH     = os.path.join(_REPO_ROOT, 'seed_kg_v4.json')

def build_entity_patterns(nodes):
    """KG 노드 → 뉴스 매칭 키워드 사전 (aliases 전체 활용). v7 노트북과 동일."""
    patterns = {}
    def add(kw, nid, name, ntype):
        kw = kw.strip()
        if len(kw) >= 2:
            patterns[kw.lower()] = (nid, name, ntype)

    _SKIP_TOKENS = {
        'strait', 'canal', 'channel', 'waterway', 'sea', 'ocean', 'gulf',
        'bay', 'port', 'passage', 'route', 'waters', 'lane', 'basin',
        'export', 'import', 'control', 'ban', 'restriction', 'crisis',
        'trade', 'supply', 'demand', 'price', 'market', 'risk', 'impact',
        'flow', 'policy', 'security', 'global', 'international',
    }
    for nid, n in nodes.items():
        ntype  = n.get("node_type", "")
        name   = n.get("name", "")
        name_en = n.get("nameEn", "")
        for alias in n.get("aliases", []):
            add(alias, nid, name, ntype)
        if name_en:
            add(name_en, nid, name, ntype)
            for tok in name_en.split():
                if len(tok) >= 3 and tok.lower() not in _SKIP_TOKENS:
                    add(tok, nid, name, ntype)
        short_id = nid.split("_", 1)[-1] if "_" in nid else nid
        if len(short_id) >= 3:
            add(short_id, nid, name, ntype)
        if name and any("\uac00" <= c <= "\ud7a3" for c in name):
            add(name, nid, name, ntype)
        if ntype == "chokepoint":
            for suffix in ["strait", "canal", "channel", "waterway"]:
                add(f"{short_id.lower()} {suffix}", nid, name, ntype)
                add(f"strait of {short_id.lower()}", nid, name, ntype)
        if ntype == "commodity_flow":
            cat = n.get("category", "")
            extra = {
                "EnergyFlow":     ["crude oil", "petroleum", "brent", "wti", "crude"],
                "GasFlow":        ["natural gas", "lng", "liquefied natural gas", "lpg"],
                "ChemicalFlow":   ["naphtha", "petrochemical", "ethylene", "propylene"],
                "MetalFlow":      ["iron ore", "steel", "coking coal", "copper"],
                "GrainFlow":      ["wheat", "corn", "soybean", "grain", "maize"],
                "FertilizerFlow": ["urea", "fertilizer", "ammonia", "adblue"],
                "SemiMaterial":   ["hydrogen fluoride", "photoresist", "semiconductor material"],
            }.get(cat, [])
            for kw in extra:
                add(kw, nid, name, ntype)
    return patterns


def match_entities(title, patterns, G, nodes):
    """기사 제목 → KG 엔티티 매칭 (긴 패턴 우선). v7 노트북과 동일."""
    title_lower = title.lower()
    matched = {}
    for pattern, (eid, ename, etype) in sorted(patterns.items(), key=lambda x: -len(x[0])):
        if re.search(r"\b" + re.escape(pattern) + r"\b", title_lower):
            if eid not in matched:
                matched[eid] = (ename, etype)
    kg_ctx_lines = []
    for eid, (ename, etype) in matched.items():
        neighbors = []
        if G is not None and eid in G:
            for _, tgt, _, ed in G.out_edges(eid, data=True, keys=True):
                neighbors.append(f"{ed.get('relation', '')}→{nodes.get(tgt, {}).get('name', tgt)}")
        kg_ctx_lines.append(f"[{etype}] {ename}({eid}): {'; '.join(neighbors[:4])}")
    return {
        "matched_entities": list(matched.keys()),
        "match_count": len(matched),
        "kg_context": "\n".join(kg_ctx_lines),
    }


def apply_kg_matching(df, patterns, G, nodes):
    """DataFrame 전체에 KG 매칭 적용. 실패 시 빈 값 유지."""
    if not patterns:
        df['kg_entities']    = ''
        df['kg_match_count'] = 0
        df['_kg_context']    = ''
        return df
    results = [match_entities(str(row.get('title', '')), patterns, G, nodes)
               for _, row in df.iterrows()]
    df['kg_entities']    = [json.dumps(m["matched_entities"]) for m in results]
    df['kg_match_count'] = [m["match_count"] for m in results]
    df['_kg_context']    = [m["kg_context"] for m in results]
    matched_count = sum(1 for m in results if m["match_count"] > 0)
    print(f"  KG 매칭: {matched_count}/{len(df)}건 (패턴 {len(patterns)}개)")
    return df


# ── KG 로드 (실패 시 빈 dict → 기존과 동일하게 작동) ──
_kg_nodes = {}
_kg_G     = None
_entity_patterns = {}

try:
    with open(KG_PATH, encoding='utf-8') as f:
        _kg_raw = json.load(f)
    _kg_nodes = _kg_raw.get("nodes", {})
    _kg_edges = _kg_raw.get("edges", [])
    _kg_G = nx.MultiDiGraph()
    for nid, ndata in _kg_nodes.items():
        _kg_G.add_node(nid, **ndata)
    for e in _kg_edges:
        _kg_G.add_edge(e["from"], e["to"],
                       **{k: v for k, v in e.items() if k not in ("from", "to")})
    _entity_patterns = build_entity_patterns(_kg_nodes)
    print(f"✅ KG 로드: {len(_kg_nodes)}노드, {len(_kg_edges)}엣지, {len(_entity_patterns)}패턴")
except FileNotFoundError:
    print(f"⚠ KG 파일 없음 ({KG_PATH}) — KG 매칭 비활성화, 기존 방식으로 분류")
except Exception as e:
    print(f"⚠ KG 로드 실패: {e} — KG 매칭 비활성화, 기존 방식으로 분류")


# ── 추적 키워드 ──
TRACKED_KEYWORDS = [
    ('호르무즈|Hormuz',                '호르무즈'),
    ('이란|Iran|IRGC',                 '이란'),
    ('트럼프|Trump',                   '트럼프'),
    ('나프타|naphtha|납사',             '나프타'),
    ('LNG',                            'LNG'),
    ('원유|crude oil|petroleum',       '원유'),
    ('유가|oil price',                 '유가'),
    ('코스피|KOSPI',                   '코스피'),
    ('환율|exchange rate|won-dollar',  '환율'),
    ('수에즈|Suez',                    '수에즈'),
    ('말라카|Malacca',                 '말라카'),
    ('운임|freight rate|tanker rate',  '운임'),
    ('셧다운|shutdown',                '셧다운'),
    ('비상경영|emergency',              '비상경영'),
    ('OECD',                           'OECD'),
    ('반도체|semiconductor',           '반도체'),
    ('석유화학|petrochemical',          '석유화학'),
    ('물가|inflation|CPI',             '물가'),
    ('홍해|Red Sea|Houthi',           '홍해'),
    ('해운|shipping|carrier',          '해운'),
]


# ══════════════════════════════════════════════════════════════
# 1. 유틸리티 함수
# ══════════════════════════════════════════════════════════════

def _make_hash(url):
    return hashlib.md5(url.encode()).hexdigest()


def load_keywords(query_file, lang='en'):
    with open(query_file, 'r', encoding='utf-8') as f:
        config = json.load(f)
    keywords, kw_source = [], {}
    for qname, qs in config['query_sets'].items():
        added = 0
        for kw in qs.get(lang, []):
            if kw not in kw_source:
                keywords.append(kw)
                kw_source[kw] = qname
                added += 1
                # 테스트 모드: 도메인별 2개만
                if _test_scale and added >= 2:
                    break
    return keywords, kw_source


def dedup_by_title(df, group_label):
    before = len(df)
    df = df.copy()
    df['_title_norm'] = df['title'].str.lower().str.strip()
    df = df.drop_duplicates(subset='_title_norm', keep='first').drop(columns=['_title_norm'])
    after = len(df)
    print(f"  [{group_label}] dedup: {before} → {after}건 (-{before-after}건)")
    return df.reset_index(drop=True)


def proportional_sample(df, max_n, min_per_kw=3, kw_col='query_keyword'):
    if len(df) <= max_n:
        return df
    kw_counts = df[kw_col].value_counts()
    kw_counts = kw_counts[kw_counts > 0]
    total     = kw_counts.sum()
    alloc = (kw_counts / total * max_n).clip(lower=min_per_kw).round().astype(int)
    diff = max_n - alloc.sum()
    if diff > 0:
        for kw in kw_counts.index:
            if diff == 0: break
            spare = int(kw_counts[kw]) - alloc[kw]
            add = min(spare, diff)
            if add > 0:
                alloc[kw] += add
                diff -= add
    elif diff < 0:
        excess = -diff
        for kw in alloc.index:
            if excess == 0: break
            reducible = alloc[kw] - min_per_kw
            reduce_by = min(reducible, excess)
            if reduce_by > 0:
                alloc[kw] -= reduce_by
                excess    -= reduce_by
    sampled = []
    for kw, n in alloc.items():
        sampled.append(df[df[kw_col] == kw].head(int(n)))
    result = pd.concat(sampled, ignore_index=True)
    print(f"  비례 샘플링: {len(df)} → {len(result)}건 (max={max_n})")
    return result


def call_llm_json(prompt, system="Return ONLY valid JSON.", max_tokens=16384, retries=3):
    text = ''
    for attempt in range(retries):
        try:
            resp = client.messages.create(
                model=MODEL,
                max_tokens=max_tokens,
                temperature=0.1,
                system=system,
                messages=[{"role": "user", "content": prompt}]
            )
            text = resp.content[0].text.strip()
            if text.startswith("```"):
                text = re.sub(r'^```\w*\n?', '', text)
                text = re.sub(r'\n?```$', '', text)
            try:
                return json.loads(text)
            except json.JSONDecodeError as je:
                if 'Extra data' in str(je):
                    decoder = json.JSONDecoder()
                    obj, _ = decoder.raw_decode(text)
                    return obj
                raise
        except (json.JSONDecodeError, Exception) as e:
            if attempt < retries - 1:
                time.sleep(1 * (attempt + 1))
                continue
            print(f"  ⚠ LLM JSON 파싱 실패 ({attempt+1}회): {e}")
            print(f"  Raw: {text[:200]}...")
            return None


SYSTEM_MSG = (
    "You are a Korean supply chain risk analyst. "
    "Classify news relevance to KOREAN maritime supply chain disruption. "
    "Return ONLY valid JSON array."
)


def build_classify_prompt(batch_articles, batch_contexts, mode='mon'):
    """분류 프롬프트 생성.
    mode='kg':  KG 엔티티 컨텍스트 활용 (KG Q1~Q7 수집 기사용)
    mode='mon': KG 없이 상세 카테고리 기준만으로 분류 (모니터링 D1~D9 수집 기사용)
    """
    items = []
    for j, (art, ctx) in enumerate(zip(batch_articles, batch_contexts)):
        title = art['title']
        lang = art.get('language', 'English')
        if ctx:
            items.append(f"{j+1}. [{lang}] {title}\n   KG entities: {ctx}")
        else:
            items.append(f"{j+1}. [{lang}] {title}")

    news_block = '\n'.join(items)

    if mode == 'mon':
        # 모니터링: KG 없이 상세 relevance 기준 + 카테고리 분류
        prompt = f"""You are classifying news headlines for KMI (Korea Maritime Institute) daily maritime supply chain monitoring.

BACKGROUND — Korea's maritime supply chain vulnerabilities:
- Korea imports by sea: crude oil 95%, LNG 99%, naphtha 100%, iron ore 100%, grain 77%
- Key chokepoints: Hormuz (70% of Korea's oil), Malacca (85% energy transit), Suez/Bab el-Mandeb, Panama, Taiwan Strait
- Key dependencies: Middle East oil 71%, Qatar LNG 19% (via Hormuz), Australia LNG 25%, China urea 97%
- Key sectors affected: refining/petrochemical (naphtha-based), steel (iron ore/coal), shipping, food/agriculture
- Key Korean actors: SK Innovation, GS Caltex, S-Oil, POSCO, HMM, KOGAS, KEPCO, Samsung, Hyundai

STEP 1 — Relevance (use the criteria below strictly):
- HIGH: Article directly reports on:
  * Chokepoint disruption, blockade, military action, or transit restriction affecting Korea-bound routes
  * Korea's import supply disruption (oil, LNG, naphtha, coal, iron ore, grain, urea, semiconductors)
  * Korean port, shipping company, or energy infrastructure directly impacted
  * Korean government policy response to supply chain crisis (SPR release, emergency measures)
  * Freight rate surge or shipping reroute that directly affects Korea trade routes
  * Vessel seizure, attack, or accident on Korea-relevant shipping lanes

- MEDIUM: Article reports on:
  * Global chokepoint tension or military posturing (not yet directly disrupting Korea)
  * International oil/LNG/commodity price movements that will impact Korea's import costs
  * Major carrier (Maersk, MSC, CMA CGM) operational changes affecting Asia routes
  * IMO regulations, maritime safety developments with potential Korea impact
  * Financial market reactions (KOSPI, KRW, VIX) to maritime/energy events
  * Regional geopolitical developments (Iran, Houthi, US Navy) with supply chain implications
  * Seafood trade disruptions or fishing industry impacts from energy price changes

- LOW: Article is about:
  * Other countries' domestic supply chain issues with no clear Korea connection
    (But if KG entities show a link to Korean sectors/ports/companies, upgrade to MEDIUM)
  * General industry trends or corporate news without supply chain disruption angle
  * Historical analysis or opinion pieces without current operational impact
  * Maritime topics unrelated to supply chain (tourism, environment, sports)

- NONE: Article has no connection to maritime supply chain monitoring

STEP 2 — Category (assign the BEST matching ONE category for HIGH and MEDIUM; leave "" for LOW/NONE):
{CAT_BLOCK_MON}

Headlines:
{news_block}

IMPORTANT: Classify based ONLY on what the headline explicitly states. Do NOT infer, assume, or add information beyond the headline text and KG context provided.

Return ONLY valid JSON array:
[{{"index": 1, "relevance": "HIGH/MEDIUM/LOW/NONE", "topic": "2-3 word topic", "category": "1_Security or empty"}}]"""

    else:
        # KG: 엔티티 컨텍스트 활용
        prompt = f"""Classify each news headline's relevance to KOREAN maritime supply chain disruption risk.
Use the KG (Knowledge Graph) entity matches to inform your judgment.
Consider Korea's dependency on maritime imports (oil 95%, LNG 99%, naphtha 100%).

STEP 1 — Relevance:
- HIGH: Direct Korea supply chain impact (chokepoint blockage, commodity shortage, shipping disruption affecting Korea)
- MEDIUM: Indirect impact (freight rate changes, route congestion, energy price spikes, regional tension)
- LOW: Related but no immediate Korea impact (global trends, other-country specific, industry news)
- NONE: Not related to Korean maritime supply chain (even if KG entities appear in non-relevant context)

STEP 2 — Category (for HIGH and MEDIUM only, leave "" for LOW/NONE):
{CAT_BLOCK_MON}

Headlines with KG context:
{news_block}

IMPORTANT: Classify based ONLY on what the headline explicitly states. Do NOT infer, assume, or add information beyond the headline text and KG context provided.

Return ONLY valid JSON array:
[{{"index": 1, "relevance": "HIGH/MEDIUM/LOW/NONE", "topic": "2-3 word topic", "category": "1_Security or empty"}}]"""

    return prompt


def classify_group(classify_df, group_label, ckpt_file, mode='mon'):
    """한 그룹의 기사를 LLM으로 분류."""
    if len(classify_df) == 0:
        print(f"\n── {group_label}: 분류할 기사 없음 ──")
        return classify_df

    print(f"\n── {group_label}: LLM 분류 시작 ({len(classify_df)}건) ──")

    # 초기값
    classify_df = classify_df.copy()
    classify_df['relevance'] = 'NONE'
    classify_df['topic'] = ''
    classify_df['category'] = ''

    # 체크포인트 로드 (중단 재개)
    resume_batch = 0
    if os.path.exists(ckpt_file):
        ckpt = pd.read_csv(ckpt_file, encoding='utf-8-sig')
        ckpt_cols = list(ckpt.columns)
        has_cat = 'category' in ckpt_cols
        restored = 0
        for _, crow in ckpt.iterrows():
            mask = classify_df['url_hash'] == crow['url_hash']
            if mask.any():
                idx = classify_df[mask].index[0]
                classify_df.loc[idx, 'relevance'] = crow['relevance']
                classify_df.loc[idx, 'topic'] = crow['topic']
                if has_cat:
                    classify_df.loc[idx, 'category'] = crow.get('category', '')
                restored += 1
        resume_batch = restored // BATCH_SIZE
        print(f"  ⚡ 체크포인트: {restored}건 복원 → 배치 {resume_batch}부터 재개")
        del ckpt

    target_idx = classify_df.index.tolist()
    total = len(target_idx)
    total_batches = (total + BATCH_SIZE - 1) // BATCH_SIZE

    llm_calls = 0
    errors = 0
    t0 = time.time()

    for batch_num in range(total_batches):
        if batch_num < resume_batch:
            continue

        start = batch_num * BATCH_SIZE
        end = min(start + BATCH_SIZE, total)
        batch_idx = target_idx[start:end]

        batch_rows = classify_df.loc[batch_idx].to_dict('records')
        batch_contexts = classify_df.loc[batch_idx, '_kg_context'].tolist()

        prompt = build_classify_prompt(batch_rows, batch_contexts, mode=mode)
        result = call_llm_json(prompt, system=SYSTEM_MSG, max_tokens=16384)
        llm_calls += 1

        if result and isinstance(result, list):
            for item in result:
                j = item.get('index', 0) - 1
                if 0 <= j < len(batch_idx):
                    rel = item.get('relevance', 'LOW').upper()
                    if rel not in ('HIGH', 'MEDIUM', 'LOW', 'NONE'):
                        rel = 'LOW'
                    classify_df.loc[batch_idx[j], 'relevance'] = rel
                    classify_df.loc[batch_idx[j], 'topic'] = item.get('topic', '')
                    # 카테고리: HIGH/MEDIUM만, 유효한 키인지 검증
                    cat = item.get('category', '')
                    if rel in ('HIGH', 'MEDIUM') and cat in CATEGORY_KEYS:
                        classify_df.loc[batch_idx[j], 'category'] = cat
                    elif rel in ('HIGH', 'MEDIUM') and cat:
                        # 부분 매칭 시도 (LLM이 숫자만 반환하는 경우 등)
                        matched_cat = [k for k in CATEGORY_KEYS if cat in k or k.startswith(cat)]
                        if matched_cat:
                            classify_df.loc[batch_idx[j], 'category'] = matched_cat[0]
                        else:
                            classify_df.loc[batch_idx[j], 'category'] = cat  # 원본 유지
        else:
            errors += 1
            for idx in batch_idx:
                classify_df.loc[idx, 'relevance'] = 'LOW'
                classify_df.loc[idx, 'topic'] = 'classification_error'

        if (batch_num + 1) % 20 == 0 or batch_num == total_batches - 1:
            elapsed = time.time() - t0
            print(f"  [{batch_num+1}/{total_batches}] LLM {llm_calls}회, 오류 {errors}건 ({elapsed:.0f}초)")

        # 매 50배치마다 체크포인트
        if llm_calls > 0 and llm_calls % 50 == 0:
            _ckpt = classify_df[['url_hash', 'title', 'relevance', 'topic', 'category']].copy()
            _ckpt.to_csv(ckpt_file, index=False, encoding='utf-8-sig')
            del _ckpt
            print(f"  💾 체크포인트 저장 ({ckpt_file})")

        time.sleep(0.3)

    # 임시 필드 정리
    if '_kg_context' in classify_df.columns:
        classify_df.drop(columns=['_kg_context'], inplace=True)

    # ── 분류 통계 ──
    elapsed = time.time() - t0
    rel_counts = classify_df['relevance'].value_counts()
    print(f"\n  [{group_label}] 분류 완료 ({elapsed:.0f}초, LLM {llm_calls}회)")
    for rel in ['HIGH', 'MEDIUM', 'LOW', 'NONE']:
        n = rel_counts.get(rel, 0)
        pct = n / len(classify_df) * 100
        print(f"    {rel:8s}: {n:5d}건 ({pct:.1f}%)")

    # KG 매칭별 분류 분포
    if 'kg_match_count' in classify_df.columns:
        print(f"    KG 매칭 유무별:")
        for label, mask in [('매칭 1+', classify_df['kg_match_count'] > 0),
                             ('매칭 0', classify_df['kg_match_count'] == 0)]:
            sub = classify_df[mask]
            if len(sub) > 0:
                dist = sub['relevance'].value_counts()
                dist_str = ', '.join(f"{r}:{dist.get(r,0)}" for r in ['HIGH','MEDIUM','LOW','NONE'])
                print(f"      {label} ({len(sub)}건): {dist_str}")

    # 카테고리별 분포 (HIGH+MEDIUM)
    hm = classify_df[classify_df['relevance'].isin(['HIGH', 'MEDIUM'])]
    if len(hm) > 0:
        print(f"\n    카테고리별 분포 (HIGH+MEDIUM {len(hm)}건):")
        cat_counts = hm['category'].value_counts()
        for cat, cnt in cat_counts.items():
            label = cat if cat else '(미분류)'
            print(f"      {label:20s} {cnt:4d}건")

    return classify_df


def save_classified(classify_df, classified_csv, daily_prefix, ckpt_file):
    """분류된 기사를 당일 CSV + 누적 CSV에 저장."""
    if len(classify_df) == 0:
        print("저장할 분류 결과 없음")
        return

    save_cols = [
        'url_hash', 'title', 'url', 'seendate', 'domain', 'language',
        'sourcecountry', 'query_keyword', 'query_group', 'collect_date',
        'kg_entities', 'kg_match_count', 'relevance', 'topic', 'category'
    ]
    save_cols = [c for c in save_cols if c in classify_df.columns]
    df_save = classify_df[save_cols].copy()

    rel_order = {'HIGH': 0, 'MEDIUM': 1, 'LOW': 2, 'NONE': 3}
    df_save['_rel_order'] = df_save['relevance'].map(rel_order)
    df_save = df_save.sort_values(['_rel_order', 'seendate'], ascending=[True, False])
    df_save = df_save.drop(columns=['_rel_order']).reset_index(drop=True)

    # ── 당일 CSV (날짜별 폴더에 저장) ──
    for cd, grp in df_save.groupby('collect_date'):
        date_tag = cd.replace('-', '')
        daily_dir = os.path.join(MONITOR_DIR, date_tag)
        os.makedirs(daily_dir, exist_ok=True)
        daily_csv = os.path.join(daily_dir, f'{daily_prefix}_daily_{date_tag}.csv')
        grp.to_csv(daily_csv, index=False, encoding='utf-8-sig')
        print(f"  당일: {daily_csv} ({len(grp)}건)")

    # ── 누적 CSV 병합 ──
    if os.path.exists(classified_csv):
        df_existing = pd.read_csv(classified_csv, encoding='utf-8-sig')
        overlap_dates = set(df_save['collect_date'].unique())
        df_keep = df_existing[~df_existing['collect_date'].isin(overlap_dates)]
        df_final = pd.concat([df_keep, df_save], ignore_index=True)
        print(f"  누적: 기존 {len(df_keep)}건 + 신규 {len(df_save)}건 = {len(df_final)}건")
    else:
        df_final = df_save.copy()
        print(f"  누적: 첫 파일 생성 ({len(df_final)}건)")

    df_final.to_csv(classified_csv, index=False, encoding='utf-8-sig')
    print(f"  → {classified_csv} ({os.path.getsize(classified_csv)//1024} KB)")

    rel_counts = df_save['relevance'].value_counts()
    high_n = rel_counts.get('HIGH', 0)
    med_n  = rel_counts.get('MEDIUM', 0)
    print(f"  분류: {len(df_save)}건 (HIGH:{high_n}, MEDIUM:{med_n})")

    if high_n > 0:
        print(f"\n  🔴 HIGH 기사:")
        for _, row in df_save[df_save['relevance'] == 'HIGH'].head(10).iterrows():
            cat_tag = f"[{row.get('category','')}] " if row.get('category','') else ''
            print(f"    {cat_tag}[{row.get('topic','')}] {row['title'][:80]}")

    if os.path.exists(ckpt_file):
        os.remove(ckpt_file)
        print(f"  🗑 체크포인트 삭제")


# ══════════════════════════════════════════════════════════════
# 2. GDELT 수집
# ══════════════════════════════════════════════════════════════

print("\n[Step 1/7] GDELT 영문 수집")
print("-" * 40)

mon_keywords, mon_kw_source = load_keywords(MON_QUERY_FILE, lang='en')
valid_keywords = [kw for kw in mon_keywords if len(kw) >= MIN_KEYWORD_LEN]
print(f"키워드: {len(valid_keywords)}개")

s_date = TARGET_DATE.strftime('%Y-%m-%d')
e_date = (TARGET_DATE + timedelta(days=1)).strftime('%Y-%m-%d')

gdelt_articles = []
seen_urls = set()
gdelt_stats, gdelt_errors = [], []
t0 = time.time()

for idx, keyword in enumerate(valid_keywords):
    qgroup = mon_kw_source[keyword]
    n_raw = n_new = 0
    err = None

    for attempt in range(MAX_RETRIES + 1):
        try:
            f = Filters(keyword=keyword, start_date=s_date, end_date=e_date,
                        num_records=MAX_RECORDS, language=LANGUAGES)
            results = gd.article_search(f)
            if results is not None and len(results) > 0:
                n_raw = len(results)
                for _, row in results.iterrows():
                    url = row.get('url', '')
                    url_hash = _make_hash(url)
                    if url_hash not in seen_urls:
                        seen_urls.add(url_hash)
                        gdelt_articles.append({
                            'url': url, 'url_hash': url_hash,
                            'title': row.get('title', ''),
                            'seendate': row.get('seendate', ''),
                            'domain': row.get('domain', ''),
                            'language': row.get('language', ''),
                            'sourcecountry': row.get('sourcecountry', ''),
                            'query_keyword': keyword,
                            'query_group': qgroup,
                            'collect_date': str(TARGET_DATE),
                        })
                        n_new += 1
            time.sleep(SLEEP_SEC)
            break
        except Exception as e:
            err = str(e)
            is_network = any(x in err for x in ['ConnectionReset','ConnectionAborted',
                                                  'RemoteDisconnected','timeout','Timeout'])
            if is_network and attempt < MAX_RETRIES:
                time.sleep(3 * (attempt + 1))
                err = None
                continue
            gdelt_errors.append({'keyword': keyword, 'error': err})
            break

    gdelt_stats.append({'keyword': keyword, 'query_group': qgroup,
                        'raw': n_raw, 'new': n_new, 'error': err or ''})

    if (idx + 1) % 20 == 0:
        print(f"  [{idx+1}/{len(valid_keywords)}] {len(gdelt_articles)}건 수집 중...")

elapsed = time.time() - t0
print(f"✅ GDELT 수집 완료: {len(gdelt_articles)}건 ({elapsed:.0f}초)")

# Raw CSV 저장
gdelt_raw_df = pd.DataFrame(gdelt_articles)
if len(gdelt_raw_df) > 0:
    raw_csv = os.path.join(DAILY_DIR, f'gdelt_mon_daily_{DATE_TAG}.csv')
    gdelt_raw_df.to_csv(raw_csv, index=False, encoding='utf-8-sig')
    print(f"  원본 저장: {raw_csv} ({len(gdelt_raw_df)}건)")


# ══════════════════════════════════════════════════════════════
# 3. 네이버 수집
# ══════════════════════════════════════════════════════════════

print("\n[Step 2/7] 네이버 한국어 수집")
print("-" * 40)

naver_articles = []

if not NAVER_CLIENT_ID or not NAVER_CLIENT_SECRET:
    print("⚠ NAVER 환경변수 미설정 — 네이버 수집 건너뜀")
    naver_raw_df = pd.DataFrame()
else:
    naver_keywords, naver_kw_source = load_keywords(MON_QUERY_FILE, lang='ko')
    print(f"키워드: {len(naver_keywords)}개")

    headers = {
        'X-Naver-Client-Id': NAVER_CLIENT_ID,
        'X-Naver-Client-Secret': NAVER_CLIENT_SECRET,
    }

    def _clean(text):
        text = re.sub(r'<[^>]+>', '', text)
        return html.unescape(text).strip()

    def _parse_naver_date(pub_date_str):
        try:
            from email.utils import parsedate_to_datetime
            return parsedate_to_datetime(pub_date_str).date()
        except Exception:
            return None

    seen_naver = set()
    naver_stats = []
    t0 = time.time()

    for idx, keyword in enumerate(naver_keywords):
        qgroup = naver_kw_source[keyword]
        n_new = 0
        start_idx = 1

        max_start = 31 if _test_scale else 1000  # 테스트: 1페이지만, 네이버 API 한계=1000
        while start_idx <= max_start:
            try:
                params = {'query': keyword, 'display': NAVER_DISPLAY,
                          'start': start_idx, 'sort': 'date'}
                resp = requests.get(NAVER_API_URL, headers=headers, params=params, timeout=10)
                resp.raise_for_status()
                data = resp.json()
                items = data.get('items', [])
                if not items:
                    break

                for item in items:
                    pub_d = _parse_naver_date(item.get('pubDate', ''))
                    if pub_d is None or pub_d < TARGET_DATE:
                        break
                    if pub_d > TARGET_DATE:
                        continue

                    url = item.get('link', item.get('originallink', ''))
                    url_hash = _make_hash(url)
                    if url_hash not in seen_naver:
                        seen_naver.add(url_hash)
                        naver_articles.append({
                            'url': url, 'url_hash': url_hash,
                            'title': _clean(item.get('title', '')),
                            'seendate': str(TARGET_DATE),
                            'domain': url.split('/')[2] if url else '',
                            'language': 'Korean',
                            'sourcecountry': 'South Korea',
                            'query_keyword': keyword,
                            'query_group': qgroup,
                            'collect_date': str(TARGET_DATE),
                        })
                        n_new += 1
                else:
                    if len(items) == NAVER_DISPLAY:
                        start_idx += NAVER_DISPLAY
                        time.sleep(NAVER_SLEEP)
                        continue
                break

                time.sleep(NAVER_SLEEP)
                break
            except Exception as e:
                print(f"  ⚠ 네이버 에러 ({keyword}): {e}")
                break

        naver_stats.append({'keyword': keyword, 'query_group': qgroup, 'new': n_new})

        if (idx + 1) % 20 == 0:
            print(f"  [{idx+1}/{len(naver_keywords)}] {len(naver_articles)}건...")

    elapsed = time.time() - t0
    print(f"✅ 네이버 수집 완료: {len(naver_articles)}건 ({elapsed:.0f}초)")

    naver_raw_df = pd.DataFrame(naver_articles)
    if len(naver_raw_df) > 0:
        naver_raw_csv = os.path.join(DAILY_DIR, f'naver_mon_daily_{DATE_TAG}.csv')
        naver_raw_df.to_csv(naver_raw_csv, index=False, encoding='utf-8-sig')
        print(f"  원본 저장: {naver_raw_csv} ({len(naver_raw_df)}건)")


# ══════════════════════════════════════════════════════════════
# 4. GDELT dedup + 비례 샘플링 + LLM 분류
# ══════════════════════════════════════════════════════════════

print("\n[Step 3/7] GDELT dedup + 샘플링 + LLM 분류")
print("-" * 40)

gdelt_classify_df = pd.DataFrame()
if len(gdelt_raw_df) > 0:
    _dedup = dedup_by_title(gdelt_raw_df, 'GDELT')
    _dedup = apply_kg_matching(_dedup, _entity_patterns, _kg_G, _kg_nodes)
    gdelt_classify_df = proportional_sample(_dedup, MAX_LLM_SAMPLE_GDELT, MIN_PER_KEYWORD)

    CLASSIFY_CKPT = os.path.join(CUMUL_DIR, 'gdelt_mon_classify_checkpoint.csv')
    CLASSIFIED_CSV = os.path.join(CUMUL_DIR, 'gdelt_mon_classified.csv')
    gdelt_classify_df = classify_group(gdelt_classify_df, 'GDELT 모니터링', CLASSIFY_CKPT, mode='mon')
    save_classified(gdelt_classify_df, CLASSIFIED_CSV, 'gdelt_mon_classified', CLASSIFY_CKPT)


# ══════════════════════════════════════════════════════════════
# 5. 네이버 dedup + 비례 샘플링 + LLM 분류
# ══════════════════════════════════════════════════════════════

print("\n[Step 4/7] 네이버 dedup + 샘플링 + LLM 분류")
print("-" * 40)

naver_classify_df = pd.DataFrame()
if len(naver_raw_df) > 0:
    _naver_dedup = dedup_by_title(naver_raw_df, '네이버')
    _naver_dedup = apply_kg_matching(_naver_dedup, _entity_patterns, _kg_G, _kg_nodes)
    naver_classify_df = proportional_sample(_naver_dedup, MAX_LLM_SAMPLE_NAVER, MIN_PER_KEYWORD)

    NAVER_CLASSIFY_CKPT = os.path.join(CUMUL_DIR, 'naver_mon_classify_checkpoint.csv')
    NAVER_CLASSIFIED_CSV = os.path.join(CUMUL_DIR, 'naver_mon_classified.csv')
    naver_classify_df = classify_group(naver_classify_df, '네이버 모니터링', NAVER_CLASSIFY_CKPT, mode='mon')
    save_classified(naver_classify_df, NAVER_CLASSIFIED_CSV, 'naver_mon_classified', NAVER_CLASSIFY_CKPT)


# ══════════════════════════════════════════════════════════════
# 6. 주간 롤링 파일 재구성 (Cell 10)
# ══════════════════════════════════════════════════════════════

print("\n[Step 5/7] 주간 롤링 재구성")
print("-" * 40)

for daily_prefix in ['gdelt_mon_classified', 'naver_mon_classified']:
    pattern = os.path.join(MONITOR_DIR, '*', f'{daily_prefix}_daily_*.csv')
    daily_files = sorted(glob.glob(pattern))
    if not daily_files:
        print(f"  {daily_prefix}: daily CSV 없음")
        continue

    week_map = {}
    for f in daily_files:
        df = pd.read_csv(f, encoding='utf-8-sig')
        if df.empty or 'collect_date' not in df.columns:
            continue
        for cd, grp in df.groupby('collect_date'):
            cd_date = pd.to_datetime(cd).date()
            sunday  = cd_date + timedelta(days=(6 - cd_date.weekday()))
            stag    = sunday.strftime('%Y%m%d')
            if stag not in week_map:
                week_map[stag] = {}
            week_map[stag][cd] = grp

    for stag in sorted(week_map.keys()):
        wk_df = pd.concat(list(week_map[stag].values()), ignore_index=True)
        weekly_csv = os.path.join(WEEKLY_DIR, f'{daily_prefix}_week_{stag}.csv')
        wk_df.to_csv(weekly_csv, index=False, encoding='utf-8-sig')
        print(f"  → {weekly_csv} ({len(wk_df)}건)")


# ══════════════════════════════════════════════════════════════
# 7. Excel 통계 리포트 (Cell 10b — _extract_report_data + _generate_xlsx)
# ══════════════════════════════════════════════════════════════

def _extract_report_data(target_date=None):
    """분류 CSV에서 리포트에 필요한 모든 데이터를 추출. dict 반환."""
    if target_date is None:
        target_date = TARGET_DATE
    td = str(target_date).replace('-', '')
    td_fmt = str(target_date)

    # CSV 로드
    daily_dir = os.path.join(MONITOR_DIR, td)
    gdelt_csv = os.path.join(daily_dir, f'gdelt_mon_classified_daily_{td}.csv')
    naver_csv = os.path.join(daily_dir, f'naver_mon_classified_daily_{td}.csv')

    dfs = []
    sources_loaded = []
    for csv_path, label in [(gdelt_csv, 'GDELT'), (naver_csv, '네이버')]:
        if os.path.exists(csv_path):
            df = pd.read_csv(csv_path, encoding='utf-8-sig')
            df['_source'] = label
            dfs.append(df)
            sources_loaded.append(f"{label}({len(df)})")
        else:
            sources_loaded.append(f"{label}(없음)")

    if not dfs:
        print("❌ 분류 CSV가 없습니다.")
        return None

    print(f"📂 로드: {', '.join(sources_loaded)}")

    df_all = pd.concat(dfs, ignore_index=True)
    df_all = df_all.drop_duplicates(subset='url_hash', keep='first').reset_index(drop=True)

    total = len(df_all)
    rel_counts = df_all['relevance'].value_counts()
    n_high = int(rel_counts.get('HIGH', 0))
    n_med  = int(rel_counts.get('MEDIUM', 0))
    n_low  = int(rel_counts.get('LOW', 0))
    n_none = int(rel_counts.get('NONE', 0))

    df_all['source_type'] = df_all['sourcecountry'].apply(
        lambda x: 'domestic' if x == 'South Korea' else 'international'
    )

    hm = df_all[df_all['relevance'].isin(['HIGH', 'MEDIUM'])].copy()
    hm_kr = hm[hm['language'] == 'Korean']
    hm_domestic = hm[hm['source_type'] == 'domestic']
    hm_intl     = hm[hm['source_type'] == 'international']

    # 카테고리별 통계
    cat_stats = {}
    for cat in CAT_ORDER:
        cat_hm = hm[hm['category'] == cat]
        h = len(cat_hm[cat_hm['relevance'] == 'HIGH'])
        m = len(cat_hm[cat_hm['relevance'] == 'MEDIUM'])
        t = h + m
        pct = t / len(hm) * 100 if len(hm) > 0 else 0
        cat_stats[cat] = {'high': h, 'med': m, 'total': t, 'pct': pct}

    # 키워드 빈도
    kw_results = []
    for pattern_str, label in TRACKED_KEYWORDS:
        count = len(hm[hm['title'].str.contains(pattern_str, case=False, na=False)])
        if count > 0:
            kw_results.append((label, count, count / len(hm) * 100 if len(hm) > 0 else 0))
    kw_results.sort(key=lambda x: -x[1])

    # 전일 데이터
    prev_date = (pd.Timestamp(td_fmt) - pd.Timedelta(days=1)).strftime('%Y%m%d')
    prev_data = None
    _prev_dfs = []
    _prev_dir = os.path.join(MONITOR_DIR, prev_date)
    for _pcsv in [os.path.join(_prev_dir, f'gdelt_mon_classified_daily_{prev_date}.csv'),
                  os.path.join(_prev_dir, f'naver_mon_classified_daily_{prev_date}.csv')]:
        if os.path.exists(_pcsv):
            _prev_dfs.append(pd.read_csv(_pcsv, encoding='utf-8-sig'))
    if _prev_dfs:
        df_prev = pd.concat(_prev_dfs, ignore_index=True)
        hm_prev = df_prev[df_prev['relevance'].isin(['HIGH', 'MEDIUM'])]
        prev_cat = {}
        for cat in CAT_ORDER:
            prev_cat[cat] = len(hm_prev[hm_prev['category'] == cat])
        prev_data = {
            'date': f"{prev_date[:4]}-{prev_date[4:6]}-{prev_date[6:]}",
            'total': len(df_prev),
            'hm_total': len(hm_prev),
            'cat': prev_cat,
        }

    return {
        'td': td, 'td_fmt': td_fmt,
        'total': total, 'n_high': n_high, 'n_med': n_med,
        'n_low': n_low, 'n_none': n_none,
        'hm': hm, 'hm_kr': hm_kr,
        'hm_domestic': hm_domestic, 'hm_intl': hm_intl,
        'cat_stats': cat_stats,
        'kw_results': kw_results,
        'prev_data': prev_data,
        'df_all': df_all,
    }


def _generate_xlsx(data):
    """openpyxl로 다중 시트 Excel 리포트 생성."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter

    td_fmt = data['td_fmt']
    hm = data['hm']
    hm_kr = data['hm_kr']
    hm_domestic, hm_intl = data['hm_domestic'], data['hm_intl']
    cat_stats = data['cat_stats']

    wb = Workbook()

    # 공통 스타일
    header_font = Font(name='Malgun Gothic', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    high_fill   = PatternFill(start_color='FCE4EC', end_color='FCE4EC', fill_type='solid')
    med_fill    = PatternFill(start_color='FFF8E1', end_color='FFF8E1', fill_type='solid')
    cell_font   = Font(name='Malgun Gothic', size=10)
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC'),
    )

    def _style_header(ws, row=1, max_col=None):
        if max_col is None:
            max_col = ws.max_column
        for col in range(1, max_col + 1):
            cell = ws.cell(row=row, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border

    def _auto_width(ws, min_width=8, max_width=60):
        for col_cells in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col_cells[0].column)
            for cell in col_cells:
                if cell.value:
                    # 한글은 대략 2배 너비
                    val = str(cell.value)
                    kr_chars = sum(1 for c in val if ord(c) > 0x1100)
                    max_len = max(max_len, len(val) + kr_chars)
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_width), max_width)

    # ── Sheet 1: 요약 (Summary) ──
    ws1 = wb.active
    ws1.title = '요약'

    ws1.cell(1, 1, f'해상 공급망 모니터링 일일 요약 — {td_fmt}').font = Font(name='Malgun Gothic', bold=True, size=14)
    ws1.merge_cells('A1:E1')

    # 통계
    ws1.cell(3, 1, '지표').font = Font(bold=True)
    ws1.cell(3, 2, '값').font = Font(bold=True)
    stats_rows = [
        ('총 수집', data['total']),
        ('HIGH', data['n_high']),
        ('MEDIUM', data['n_med']),
        ('LOW', data['n_low']),
        ('NONE', data['n_none']),
        ('HIGH+MEDIUM', len(hm)),
        ('한국어 (H+M)', len(hm_kr)),
    ]
    for i, (k, v) in enumerate(stats_rows, 4):
        ws1.cell(i, 1, k).font = cell_font
        ws1.cell(i, 2, v).font = cell_font

    # 카테고리 분포
    cat_start = len(stats_rows) + 6
    ws1.cell(cat_start, 1, '카테고리별 분포').font = Font(name='Malgun Gothic', bold=True, size=12)
    ws1.merge_cells(f'A{cat_start}:E{cat_start}')

    h_row = cat_start + 1
    for j, h in enumerate(['카테고리', 'HIGH', 'MEDIUM', '합계', '비중(%)'], 1):
        ws1.cell(h_row, j, h)
    _style_header(ws1, row=h_row, max_col=5)

    for i, cat in enumerate(CAT_ORDER, h_row + 1):
        s = cat_stats[cat]
        ws1.cell(i, 1, CAT_KR[cat]).font = cell_font
        ws1.cell(i, 2, s['high']).font = cell_font
        ws1.cell(i, 3, s['med']).font = cell_font
        ws1.cell(i, 4, s['total']).font = cell_font
        ws1.cell(i, 5, round(s['pct'], 1)).font = cell_font
        for col in range(1, 6):
            ws1.cell(i, col).border = thin_border

    _auto_width(ws1)

    # ── Sheet 2: 전체 기사 목록 (HIGH+MEDIUM) ──
    ws2 = wb.create_sheet('기사목록_HIGH_MED')
    cols = ['relevance', 'category', 'language', 'title', 'topic', 'domain', 'url']
    col_kr = ['관련도', '카테고리', '언어', '제목', '토픽', '도메인', 'URL']
    for j, h in enumerate(col_kr, 1):
        ws2.cell(1, j, h)
    _style_header(ws2, max_col=len(cols))

    for i, (_, row) in enumerate(hm.iterrows(), 2):
        for j, col in enumerate(cols, 1):
            val = row.get(col, '')
            ws2.cell(i, j, val if pd.notna(val) else '').font = cell_font
            ws2.cell(i, j).border = thin_border
        # 조건부 색상
        if row.get('relevance') == 'HIGH':
            for j in range(1, len(cols)+1):
                ws2.cell(i, j).fill = high_fill
        elif row.get('relevance') == 'MEDIUM':
            for j in range(1, len(cols)+1):
                ws2.cell(i, j).fill = med_fill

    # 자동 필터
    ws2.auto_filter.ref = f'A1:{get_column_letter(len(cols))}{len(hm)+1}'
    _auto_width(ws2)

    # ── Sheet 3: 국내 기사 ──
    ws3 = wb.create_sheet('국내기사')
    for j, h in enumerate(col_kr, 1):
        ws3.cell(1, j, h)
    _style_header(ws3, max_col=len(cols))

    for i, (_, row) in enumerate(hm_domestic.iterrows(), 2):
        for j, col in enumerate(cols, 1):
            val = row.get(col, '')
            ws3.cell(i, j, val if pd.notna(val) else '').font = cell_font
            ws3.cell(i, j).border = thin_border
        if row.get('relevance') == 'HIGH':
            for j in range(1, len(cols)+1):
                ws3.cell(i, j).fill = high_fill

    ws3.auto_filter.ref = f'A1:{get_column_letter(len(cols))}{len(hm_domestic)+1}'
    _auto_width(ws3)

    # ── Sheet 4: 키워드 빈도 ──
    ws4 = wb.create_sheet('키워드빈도')
    for j, h in enumerate(['키워드', '언급 건수', '비중(%)'], 1):
        ws4.cell(1, j, h)
    _style_header(ws4, max_col=3)

    for i, (label, count, pct) in enumerate(data['kw_results'], 2):
        ws4.cell(i, 1, label).font = cell_font
        ws4.cell(i, 2, count).font = cell_font
        ws4.cell(i, 3, round(pct, 1)).font = cell_font
        for j in range(1, 4):
            ws4.cell(i, j).border = thin_border

    _auto_width(ws4)

    # ── Sheet 5: 전일 대비 (있을 경우) ──
    prev = data['prev_data']
    if prev:
        ws5 = wb.create_sheet('전일대비')
        for j, h in enumerate(['카테고리', f"전일({prev['date']})", f"당일({td_fmt})", '변화', '방향'], 1):
            ws5.cell(1, j, h)
        _style_header(ws5, max_col=5)

        row_idx = 2
        for cat in CAT_ORDER:
            p = prev['cat'][cat]; c = cat_stats[cat]['total']
            if p > 0 or c > 0:
                diff = c - p
                arrow = "↑" if diff > 0 else ("↓" if diff < 0 else "→")
                ws5.cell(row_idx, 1, CAT_KR[cat]).font = cell_font
                ws5.cell(row_idx, 2, p).font = cell_font
                ws5.cell(row_idx, 3, c).font = cell_font
                ws5.cell(row_idx, 4, diff).font = cell_font
                ws5.cell(row_idx, 5, arrow).font = cell_font
                for j in range(1, 6):
                    ws5.cell(row_idx, j).border = thin_border
                row_idx += 1

        _auto_width(ws5)

    out_path = os.path.join(DAILY_DIR, f'daily_report_{DATE_TAG}.xlsx')
    wb.save(out_path)
    return out_path

print("\n[Step 6/7] Excel 통계 리포트 생성")
print("-" * 40)

xlsx_data = _extract_report_data(TARGET_DATE)
if xlsx_data is not None:
    try:
        xlsx_path = _generate_xlsx(xlsx_data)
        print(f"✅ Excel 저장: {xlsx_path}")
    except Exception as e:
        print(f"  ⚠ Excel 생성 실패: {e}")
else:
    print("  분류 CSV 없음 — Excel 생성 건너뜀")


# ══════════════════════════════════════════════════════════════
# 8. LLM 일별 리포트 생성 (Cell 12)
# ══════════════════════════════════════════════════════════════

print("\n[Step 7/7] LLM 일별 리포트 + HTML 뷰어 생성")
print("-" * 40)

LLM_REPORT_MODEL = "claude-sonnet-4-6"
REPORT_SYSTEM = """당신은 KMI(한국해양수산개발원) 해상 공급망 위기 모니터링 시스템의 전문 분석가입니다.
현재 핵심 사태: 2026년 호르무즈 해협 봉쇄 위기 (역사상 최초 실제 봉쇄)
분석 관점: 글로벌 해상 공급망 교란 → 한국 경제·산업·소비자에 미치는 영향
원칙: 기사 제목에 명시된 사실만 서술. 기사에 없는 내용을 절대 추가·추론·상상하지 말 것. 과장 금지. 불확실하면 '~로 보임', '~가능성' 명시.
출력: 반드시 유효한 JSON만 출력. 설명 텍스트나 마크다운 코드블록 없이 JSON 객체만."""



def _build_report_prompt(data):
    """LLM 호출용 통합 프롬프트 (categories + flow + changes 한 번에)"""
    td_fmt   = data['td_fmt']
    hm       = data['hm']
    hm_intl  = data['hm_intl']
    hm_dom   = data['hm_domestic']

    # ── 카테고리별 기사 블록 ──
    active_cats    = []
    articles_block = ""
    for cat in CAT_ORDER:
        cat_hm = hm[hm['category'] == cat]
        if len(cat_hm) == 0:
            continue
        active_cats.append(cat)
        intl_rows = cat_hm[cat_hm['source_type'] == 'international'].head(10)
        dom_rows  = cat_hm[cat_hm['source_type'] == 'domestic'].head(10)
        h = len(cat_hm[cat_hm['relevance'] == 'HIGH'])
        m = len(cat_hm[cat_hm['relevance'] == 'MEDIUM'])
        articles_block += f"\n[{CAT_KR[cat]}] HIGH:{h} MEDIUM:{m}\n"
        if len(intl_rows) > 0:
            articles_block += f"  해외({len(intl_rows)}건):\n"
            for _, r in intl_rows.iterrows():
                articles_block += f"    - {r['title']}\n"
        if len(dom_rows) > 0:
            articles_block += f"  국내({len(dom_rows)}건):\n"
            for _, r in dom_rows.iterrows():
                articles_block += f"    - {r['title']}\n"

    cat_keys = ', '.join(f'"{c}"' for c in active_cats)

    prompt = f"""분석 대상일: {td_fmt}
HIGH+MEDIUM 기사: 총 {len(hm)}건 (해외 {len(hm_intl)}건 / 국내 {len(hm_dom)}건)

{articles_block}

아래 JSON 형식으로 분석하세요. 활성 카테고리 키: {cat_keys}

{{
  "executive_summary": "오늘의 핵심 동향 3~4문장. 가장 중요한 변화·위험 신호 중심. 구체적 수치·고유명사 활용.",

  "categories": {{
    "<카테고리코드>": {{
      "overseas": "해외 상황 2~4문장. 국제 기사 기반, 지금 무슨 일이 벌어지고 있는지. 없으면 빈 문자열.",
      "korea_impact": "국내 영향 2~4문장. 국내 기사 기반 한국 산업·경제 파급. 국내 기사 없으면 해외 상황의 한국 파급 가능성 서술."
    }}
  }},

  "flow": {{
    "triggers": {{
      "route":     {{ "overseas": "항로·초크포인트 차단·위협 해외 상황 2~4문장. 없으면 빈 문자열.", "korea_impact": "한국 수입 항로 위협 파급 2~4문장." }},
      "source":    {{ "overseas": "공급원 교란(수출금지·제재·감산) 해외 상황 2~4문장. 없으면 빈 문자열.", "korea_impact": "한국 원자재·에너지 공급원 파급 2~4문장." }},
      "logistics": {{ "overseas": "운임·항만·선박 교란 해외 상황 2~4문장. 없으면 빈 문자열.", "korea_impact": "한국 해운·수출입 물류 파급 2~4문장." }}
    }},
    "domestic_impact": {{
      "경제·금융":  "국내 거시경제·금융시장 영향 2~3문장. 관련 기사 없으면 빈 문자열.",
      "정유·화학":  "정유·석유화학 산업 영향 2~3문장. 없으면 빈 문자열.",
      "해운·물류":  "국내 해운사·물류 업계 영향 2~3문장. 없으면 빈 문자열.",
      "식품·농업":  "식품·농업·수산 영향 2~3문장. 없으면 빈 문자열.",
      "기타산업":   "그 외 주목할 산업 영향 1~2문장. 없으면 빈 문자열."
    }}
  }},

  "changes": {{
    "new":       ["어제 없었다가 오늘 새로 등장한 이슈 (없으면 빈 배열)"],
    "escalated": ["어제보다 심화·확대된 이슈 (없으면 빈 배열)"],
    "resolved":  ["어제 있었으나 오늘 소멸·완화된 이슈 (없으면 빈 배열)"]
  }}
}}

분석 원칙:
- 기사 제목에 명시된 사실만 서술. 기사에 없는 내용을 추가하거나 추론·상상하지 말 것
- 과장 없이 현황 중심. 불확실한 내용은 '~로 보임', '~가능성' 명시
- 한국 경제·산업 시사점 반드시 포함
- JSON 외 다른 텍스트 출력 금지"""

    return prompt, active_cats

# ── 데이터 추출 ──
report_data = _extract_report_data(TARGET_DATE)

if report_data is None:
    print("  분류 결과 없음 — LLM 리포트 건너뜀")
elif len(report_data['hm']) == 0:
    print("  HIGH+MEDIUM 기사 없음 — LLM 리포트 건너뜀")
else:
    hm = report_data['hm']
    print(f"\n🤖 LLM 리포트 생성 중 ({report_data['td_fmt']}, HIGH+MEDIUM {len(hm)}건)...")

    report_prompt, active_cats = _build_report_prompt(report_data)
    print(f"  프롬프트 길이: {len(report_prompt)}자")

    report_client = anthropic.Anthropic()
    resp = report_client.messages.create(
        model=LLM_REPORT_MODEL,
        max_tokens=16384,
        temperature=0.1,
        system=REPORT_SYSTEM,
        messages=[{"role": "user", "content": report_prompt}]
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = re.sub(r'^```\w*\n?', '', raw)
        raw = re.sub(r'\n?```$', '', raw)

    try:
        report_json = json.loads(raw)
    except Exception as e1:
        print(f"⚠ JSON 1차 파싱 실패: {e1}")
        print(f"  raw 앞 300자: {raw[:300]}")
        try:
            from json_repair import repair_json
            report_json = json.loads(repair_json(raw))
            print("✅ json-repair로 복구 성공")
        except Exception as e2:
            print(f"⚠ json-repair 실패: {e2}")
            decoder = json.JSONDecoder()
            report_json, _ = decoder.raw_decode(raw)

    # ── JSON 저장 (v3 notebook 동일 flat 구조) ──
    json_path = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.json')
    json_data = {
        'date':     str(TARGET_DATE),
        'date_key': DATE_TAG,
        'n_total':  report_data['total'],
        'n_high':   report_data['n_high'],
        'n_med':    report_data['n_med'],
        'n_low':    report_data['n_low'],
        'n_none':   report_data['n_none'],
        'n_intl':   int(len(report_data['hm_intl'])),
        'n_dom':    int(len(report_data['hm_domestic'])),
        'llm_result': report_json,
    }
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)
    print(f"  JSON 저장: {json_path}")

    # ── Markdown (v3 notebook 동일 구조) ──
    md_path = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.md')
    lines = []
    L = lines.append
    L(f"# 해상 공급망 모니터링 일일 분석 — {TARGET_DATE}")
    L(f"")
    L(f"> 생성 모델: {LLM_REPORT_MODEL}  |  분석 기준: HIGH+MEDIUM {len(report_data['hm'])}건 "
      f"(해외 {len(report_data['hm_intl'])}건 / 국내 {len(report_data['hm_domestic'])}건)")
    L(f"")
    L("---")
    L("")
    L("## 📌 오늘의 핵심")
    L("")
    L(report_json.get("executive_summary", ""))
    L("")
    L("---")
    L("")
    L("## 카테고리별 분석")
    L("")
    cats_out = report_json.get("categories", {})
    for cat in CAT_ORDER:
        if cat not in cats_out: continue
        cat_data = cats_out[cat]
        cat_kr   = CAT_KR.get(cat, cat)
        s        = report_data['cat_stats'].get(cat, {})
        h, m     = s.get('high', 0), s.get('med', 0)
        if h + m == 0: continue
        L(f"### {cat_kr}  _(HIGH {h} / MEDIUM {m})_")
        L("")
        overseas     = cat_data.get("overseas", "").strip()
        korea_impact = cat_data.get("korea_impact", "").strip()
        if overseas:
            L("**🌐 해외 상황**"); L(""); L(overseas); L("")
        if korea_impact:
            L("**🇰🇷 국내 영향**"); L(""); L(korea_impact); L("")
    L("---"); L("")
    L("## 통계 요약"); L("")
    L("| 구분 | 건수 | 비중 |"); L("|------|-----:|-----:|")
    for label, cnt in [('총 수집', report_data['total']), ('HIGH', report_data['n_high']),
                       ('MEDIUM', report_data['n_med']), ('LOW', report_data['n_low']),
                       ('NONE', report_data['n_none'])]:
        L(f"| {label} | {cnt:,} | {cnt/report_data['total']*100:.1f}% |")
    L("")
    L("| 카테고리 | HIGH | MEDIUM | 합계 |"); L("|----------|-----:|-------:|-----:|")
    for cat in CAT_ORDER:
        s = report_data['cat_stats'].get(cat, {})
        if s.get('total', 0) > 0:
            L(f"| {CAT_KR[cat]} | {s['high']} | {s['med']} | {s['total']} |")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"  MD 저장: {md_path}")

    # ── DOCX (v3 notebook 동일 구조) ──
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3.0); sec.right_margin  = Cm(3.0)
        p = doc.add_heading(f'해상 공급망 모니터링 일일 분석 — {TARGET_DATE}', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f'생성 모델: {LLM_REPORT_MODEL}  |  HIGH+MEDIUM {len(report_data["hm"])}건 '
            f'(해외 {len(report_data["hm_intl"])}건 / 국내 {len(report_data["hm_domestic"])}건)'
        )
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        doc.add_paragraph()
        doc.add_heading('오늘의 핵심', level=2)
        doc.add_paragraph(report_json.get('executive_summary', ''))
        doc.add_paragraph()
        doc.add_heading('카테고리별 분석', level=2)
        for cat in CAT_ORDER:
            if cat not in cats_out: continue
            cat_data = cats_out[cat]
            cat_kr   = CAT_KR.get(cat, cat)
            s        = report_data['cat_stats'].get(cat, {})
            h, m_cnt = s.get('high', 0), s.get('med', 0)
            if h + m_cnt == 0: continue
            doc.add_heading(f'{cat_kr}  (HIGH {h} / MEDIUM {m_cnt})', level=3)
            overseas     = cat_data.get('overseas', '').strip()
            korea_impact = cat_data.get('korea_impact', '').strip()
            if overseas:
                p = doc.add_paragraph(); p.add_run('🌐 해외 상황').bold = True
                doc.add_paragraph(overseas)
            if korea_impact:
                p = doc.add_paragraph(); p.add_run('🇰🇷 국내 영향').bold = True
                doc.add_paragraph(korea_impact)
        doc.add_page_break()
        doc.add_heading('통계 요약', level=2)
        tbl = doc.add_table(rows=6, cols=3); tbl.style = 'Table Grid'
        for j, h_txt in enumerate(['구분', '건수', '비중']):
            tbl.rows[0].cells[j].text = h_txt
            tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        for i_r, (label, cnt) in enumerate([
            ('총 수집', report_data['total']), ('HIGH',   report_data['n_high']),
            ('MEDIUM',  report_data['n_med']),  ('LOW',    report_data['n_low']),
            ('NONE',    report_data['n_none'])
        ]):
            tbl.rows[i_r+1].cells[0].text = label
            tbl.rows[i_r+1].cells[1].text = str(cnt)
            tbl.rows[i_r+1].cells[2].text = f"{cnt/report_data['total']*100:.1f}%"
        docx_path = os.path.join(DAILY_DIR, f'daily_report_llm_{DATE_TAG}.docx')
        doc.save(docx_path)
        print(f"  DOCX 저장: {docx_path}")
    except Exception as e:
        print(f"  ⚠ DOCX 생성 실패: {e}")

    print(f"✅ LLM 리포트 생성 완료")


# ══════════════════════════════════════════════════════════════
# 9. daily_brief.html 생성 (Cell 13)
# ══════════════════════════════════════════════════════════════

print("\n[HTML 뷰어 생성]")
print("-" * 40)

pattern = os.path.join(MONITOR_DIR, '????????', 'daily_report_llm_????????.json')
json_files = sorted(glob.glob(pattern), reverse=True)

if json_files:
    days = []
    for jf in json_files[:30]:  # 최근 30일
        try:
            with open(jf, encoding='utf-8') as f:
                days.append(json.load(f))
        except Exception:
            pass

    if days:
        # HTML 뷰어는 별도 스크립트로 분리 (노트북 Cell 13 로직)
        import subprocess
        result = subprocess.run(
            ['python3', 'scripts/build_viewer.py'],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            print("✅ daily_brief.html 생성 완료")
        else:
            print(f"  ⚠ 뷰어 생성 실패: {result.stderr[:200]}")
else:
    print("  JSON 없음 — 뷰어 생성 건너뜀")

print(f"\n{'='*60}")
print(f"  ✅ 일일 파이프라인 완료: {TARGET_DATE}")
print(f"  산출물 위치: monitoring/{DATE_TAG}/")
print(f"{'='*60}")
